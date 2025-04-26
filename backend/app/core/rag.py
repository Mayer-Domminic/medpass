from typing import List

from app.models import (
    Faculty,
    Document,
    DocumentChunk,
)

from ..core.database import get_db
from app.services.gemini_service import embed_texts, embed_text

from sqlalchemy import text, func
from sqlalchemy import select, desc

import pypdf
import docx
import os
from langchain.text_splitter import RecursiveCharacterTextSplitter

def extract_text_pdf(filepath: str) -> str:
    
    with open(filepath, "rb") as file:
        pdf_reader = pypdf.PdfReader(file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
    return text

def extract_text_docx(filepath: str) -> str:
    
    doc = docx.Document(filepath)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text

def extract_text(filepath: str) -> str:
    
    _, ext = os.path.splitext(filepath)
    if ext.lower() == ".pdf":
        return extract_text_pdf(filepath)
    elif ext in [".docx", ".doc"]:
        return extract_text_docx(filepath)
    else:
        raise ValueError(f"Unsupported file type: {ext}")
    
def split_text(text: str) -> List[str]:
    # Split the text into chunks of 1000 characters with a chunk overlap of 200 characters
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200,
        length_function=len
    )
    chunks = text_splitter.split_text(text)
    return chunks

def generate_document_embeddings(chunks):
    # Generate embeddings for the chunks
    embeddings = embed_texts(chunks)
    return embeddings


def ingest_document(db, file_path):
    
    try:
        text = extract_text(file_path)
    except ValueError as e:
        print(f"Error extracting text: {e}")
        return
    except Exception as e:
        print(f"Unexpected error: {e}")
        return
    
    chunks = split_text(text)
    
    embeddings = generate_document_embeddings(chunks)
    
    filename = os.path.basename(file_path)
    
    db = next(get_db())
    try:
        doc = Document(
            title=filename,
            author="System",  
            facultyid=1,  
        )
        db.add(doc)
        db.flush()
        
        for i, (text_chunk, embedding) in enumerate(zip(chunks, embeddings)):
            chunk = DocumentChunk(
                documentid=doc.documentid,
                chunkindex=i,
                content=text_chunk,
                embedding=embedding
            )
            db.add(chunk)
        
        db.commit()
        
        print(f"Document {filename} ingested successfully with ID: {doc.documentid}")
        
        return doc.documentid
      
    except Exception as e:
        db.rollback()
        print(f"Error processing file: {file_path}, {e}")
        return None
    
    finally:
        db.close()
        
def ingest_document_directory(directory_path: str):
    
    doc_ids = []
    try:
        db = next(get_db())
        for filename in os.listdir(directory_path):
            file_path = os.path.join(directory_path, filename)
            
            if os.path.isfile(file_path):
                print(f"Processing file: {file_path}")
                doc_id = ingest_document(db, file_path)
                if doc_id:
                    doc_ids.append(doc_id)
                    
        return doc_ids
    except Exception as e:
        print(f"Error processing directory: {directory_path}, {e}")
        return None
    finally:
        db.close()
        
def search_documents(query: str, limit: int = 5, faculty_id: int = None, similiarity_threshold: float = 0.5):
    
    query_embedding = embed_text(query)
    
    try:
        db = next(get_db())

        try:
            
            context = (
                select(
                    DocumentChunk.documentchunkid,
                    DocumentChunk.content,
                    Document.title,
                    Document.author,
                    
                    # Calculating vector similarity (cosine sim)
                    (1 - DocumentChunk.embedding.cosine_distance(query_embedding)).label("similarity")
                ).join(
                    Document, DocumentChunk.documentid == Document.documentid
                ).order_by(
                    desc("similarity")
                ).limit(limit)
            )
                
            if faculty_id:
                context = context.where(Document.facultyid == faculty_id)
            
            # Filtering by min similarity threshold 
            context = context.where(
                (1 - DocumentChunk.embedding.cosine_distance(query_embedding)) >= similiarity_threshold
            )
            
            results = db.execute(context).fetchall()
            
            formatted_results = [
                {
                    "documentchunkid": result[0],
                    "content": result[1],
                    "title": result[2],
                    "author": result[3],
                    "similarity": result[4]
                }
                for result in results
            ]
            
            return formatted_results
                
        except Exception as e:
            print(f"Error searching documents: {e}")
            return []
            
        finally:
            db.close()
                
            
    except Exception as e:
        print(f"Error connecting to database: {e}")
        return None
    
    