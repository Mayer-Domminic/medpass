from typing import List, Tuple

from app.models import (
    Question,
    ChatMessage,
    ChatContext,
    ChatMessageContext,
    Faculty,
    Document,
    DocumentChunk,
    Student
)

from ..core.database import get_db, get_question_with_details, get_chat_context, get_chat_message
from app.services.gemini_service import embed_texts, embed_text

from sqlalchemy import text, func
from sqlalchemy import select, desc

import pypdf
import docx
import os
from langchain.text_splitter import RecursiveCharacterTextSplitter
from app.services.s3_service import s3, BUCKET


#Pre Processing Step Before Embedding Converts a question to a string takes in a dictionary found in get_question_with_details
def convert_question_to_text(question_response: dict) -> str:
    
    question = question_response['Question']
    option = question_response['Options']
    content_areas = question_response['ContentAreas']
    
    prompt = question['Prompt']
    
    # This is for a way to format our options for context in a A, B, C, D format (chr 65 is a + 1 is each letter after)
    option_lines = [
        f"{chr(65 + i)}. {opt['OptionDescription']}" for i, opt in enumerate(option)
    ]
    
    difficulty = question.get('QuestionDifficulty', 'Unknown')
    content_names = [ca["ContentName"] for ca in content_areas]
    # If multiple content areas combine them into a single string
    content_area_line = ', '.join(content_names) if content_names else "Uncatergorized"
    
    context_lines = [
        f"Difficulty: {difficulty}",
        f"Content Areas: {content_area_line}",
        ""
    ]
    
    text = "\n".join(context_lines + [prompt] + option_lines)
    
    return text

def generate_question_embedding(question_id, db):
    question_data = get_question_with_details(question_id, db)
    if not question_data:
        return None
    question_text = convert_question_to_text(question_data)
    # Use Gemini embedding
    embedding = embed_text(question_text)
    question = db.query(Question).filter(Question.questionid == question_id).first()
    if question:
        question.embedding = embedding
        db.commit()

def generate_chatcontext_text(context: dict) -> str:
    
    title = context.get('Title', '')
    content = context.get('Content', '')
    
    # Make sure when creating metadata it has topic and tags also
    metadata = context.get('Metadata', {})
    topic = metadata.get('topic', '')
    tags = metadata.get('tags', [])
    tags_line = ', '.join(tags) if tags else "Uncatergorized"
    
    context_lines = [
        f"Title: {title}",
        f"Content: {content}",
        f"Topic: {topic}",
        f"Tags: {tags_line}",
        ""
    ]
    
    text = '\n'.join(context_lines)
    
    return text

def generate_messagecontext_text(message: dict) -> str:
    
    sender = message.get('SenderType', '')
    content = message.get('Content', '')
    
    metadata = message.get('Metadata', {})
    model = metadata.get('model', '')
    temperature = metadata.get('temperature', '')
    client = metadata.get('client', '')
    
    # Joins all the metadata together into one line as a string
    #Only attaches model and temperature if they are from a bot and not a user
    
    metadata_text = []
    
    if sender == 'bot':
        if model:
            metadata_text.append(f"Model: {model}")
        if temperature:
            metadata_text.append(f"Temperature: {temperature}")
    if client:
        metadata_text.append(f"Client: {client}")
        
    text = '\n'.join(metadata_text)
    
    context_lines = [
        f"Sender: {sender}",
        f"Content: {content}",
        text,
        ""
    ]
    
    text = '\n'.join(context_lines)
    
    return text

# ----------- WARNING ---------------------
#These Chat Embedding Functions are used to populate on ingest of the chat context and message in the fake script
# DO NOT USE THESE IN MAIN FUNCTIONALITY
# Will be removed in the future, but for now they are used to populate the database with embeddings

def generate_chatcontext_embedding_SCRIPT(chatcontext_id, db):
    context_data = get_chat_context(chatcontext_id, db) 
    if not context_data:
        return None
    context_text = generate_chatcontext_text(context_data)
    # Use Gemini embedding
    embedding = embed_text(context_text)
    context = db.query(ChatContext).filter(ChatContext.contextid == chatcontext_id).first()
    if context:
        context.embedding = embedding
        db.commit()
    
        
def generate_chatmessage_embedding_SCRIPT(chatmessage_id, db):
    message_data = get_chat_message(chatmessage_id, db)
    if not message_data:
        return None
    message_text = generate_messagecontext_text(message_data)
    # Use Gemini embedding
    embedding = embed_text(message_text)
    message = db.query(ChatMessage).filter(ChatMessage.messageid == chatmessage_id).first()
    if message:
        message.embedding = embedding
        db.commit()

# ----------- WARNING ---------------------

def extract_text_pdf(filepath: str) -> str:
    
    with open(filepath, "rb") as file:
        pdf_reader = pypdf.PdfReader(file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
    return text

def extract_text_docx(filepath: str) -> str:
    
    doc = docx.Document(filepath)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text

def extract_text(filepath: str) -> str:
    
    _, ext = os.path.splitext(filepath)
    if ext.lower() == ".pdf":
        return extract_text_pdf(filepath)
    elif ext in [".docx", ".doc"]:
        return extract_text_docx(filepath)
    else:
        raise ValueError(f"Unsupported file type: {ext}")
    
def split_text(text: str) -> List[str]:
    # Split the text into chunks of 1000 characters with a chunk overlap of 200 characters
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200,
        length_function=len
    )
    chunks = text_splitter.split_text(text)
    return chunks

# Nearly identical to the os extract but is support for S3 file objects
def extract_text_pdf_from_s3(file_obj) -> str:
    
    pdf_reader = pypdf.PdfReader(file_obj)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text() + "\n"
    return text

def extract_text_docx_from_s3(file_obj) -> str:
    doc = docx.Document(file_obj)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text

def extract_text_from_s3(file_obj, filename: str) :

    _, ext = os.path.splitext(filename)
    if ext.lower() == ".pdf":
        return extract_text_pdf_from_s3(file_obj)
    elif ext.lower() in [".docx", ".doc"]:
        return extract_text_docx_from_s3(file_obj)
    else:
        raise ValueError(f"Unsupported file type: {ext}")

def generate_document_embeddings(chunks):
    # Generate embeddings for the chunks
    embeddings = embed_texts(chunks)
    return embeddings


def ingest_document(db, file_path):
    
    try:
        text = extract_text(file_path)
    except ValueError as e:
        print(f"Error extracting text: {e}")
        return
    except Exception as e:
        print(f"Unexpected error: {e}")
        return
    
    chunks = split_text(text)
    
    embeddings = generate_document_embeddings(chunks)
    
    filename = os.path.basename(file_path)
    
    db = next(get_db())
    try:
        doc = Document(
            title=filename,
            author="System",  
            facultyid=1,  
        )
        db.add(doc)
        db.flush()
        
        for i, (text_chunk, embedding) in enumerate(zip(chunks, embeddings)):
            chunk = DocumentChunk(
                documentid=doc.documentid,
                chunkindex=i,
                content=text_chunk,
                embedding=embedding
            )
            db.add(chunk)
        
        db.commit()
        
        print(f"Document {filename} ingested successfully with ID: {doc.documentid}")
        
        return doc.documentid
      
    except Exception as e:
        db.rollback()
        print(f"Error processing file: {file_path}, {e}")
        return None
    
    finally:
        db.close()
        
def ingest_document_directory(directory_path: str):
    
    doc_ids = []
    try:
        db = next(get_db())
        for filename in os.listdir(directory_path):
            file_path = os.path.join(directory_path, filename)
            
            if os.path.isfile(file_path):
                print(f"Processing file: {file_path}")
                doc_id = ingest_document(db, file_path)
                if doc_id:
                    doc_ids.append(doc_id)
                    
        return doc_ids
    except Exception as e:
        print(f"Error processing directory: {directory_path}, {e}")
        return None
    finally:
        db.close()
        
def ingest_document_from_s3(db, bucket_name, key, logininfo_id):
    try:
        response = s3.get_object(Bucket=bucket_name, Key=key)
        content = response['Body'].read()
        
        filename = key.split('/')[-1]
        
        import io
        
        file_obj = io.BytesIO(content)
        
        text = extract_text_from_s3(file_obj, filename)
        
        chunks = split_text(text)
        embeddings = generate_document_embeddings(chunks)
        
        student = db.query(Student).filter(Student.logininfoid == logininfo_id).first()
        
        if not student:
            return None
        
        doc = Document(
            title=filename,
            author=student.firstname + " " + student.lastname,
            studentid=student.studentid,
            s3_key=key
        )
        
        db.add(doc)
        db.flush()
        
        for i, (text_chunk, embedding) in enumerate(zip(chunks, embeddings)):
            chunk = DocumentChunk(
                documentid=doc.documentid,
                chunkindex=i,
                content=text_chunk,
                embedding=embedding
            )
            db.add(chunk)
        db.commit()
        
        print(f"Document {filename} ingested successfully with ID: {doc.documentid}")
        
        return doc.documentid
        
    except s3.exceptions.NoSuchKey:
        print(f"File not found in S3: {key}")
        return None
        
    except Exception as e:
        print(f"Unexpected error with S3 file: {key}, {e}")
        return None
        
def search_documents(query: str, limit: int = 5, faculty_id: int = None, similiarity_threshold: float = 0.5):
    
    query_embedding = embed_text(query)
    
    try:
        db = next(get_db())

        try:
            
            context = (
                select(
                    DocumentChunk.documentchunkid,
                    DocumentChunk.content,
                    Document.title,
                    Document.author,
                    
                    # Calculating vector similarity (cosine sim)
                    (1 - DocumentChunk.embedding.cosine_distance(query_embedding)).label("similarity")
                ).join(
                    Document, DocumentChunk.documentid == Document.documentid
                ).order_by(
                    desc("similarity")
                ).limit(limit)
            )
                
            if faculty_id:
                context = context.where(Document.facultyid == faculty_id)
            
            # Filtering by min similarity threshold 
            context = context.where(
                (1 - DocumentChunk.embedding.cosine_distance(query_embedding)) >= similiarity_threshold
            )
            
            results = db.execute(context).fetchall()
            
            formatted_results = [
                {
                    "documentchunkid": result[0],
                    "content": result[1],
                    "title": result[2],
                    "author": result[3],
                    "similarity": result[4]
                }
                for result in results
            ]
            
            return formatted_results
                
        except Exception as e:
            print(f"Error searching documents: {e}")
            return []
            
        finally:
            db.close()
                
            
    except Exception as e:
        print(f"Error connecting to database: {e}")
        return None
    
# Chat Embedding
def generate_chat_context_embedding(db, chatcontext_id):
    
    context = db.query(ChatContext).filter(ChatContext.contextid == chatcontext_id).first()
    
    if not context:
        return None
    
    # Format Context Text Package by Packaging  Title, Content, and Metadata
    
    context_parts = [context.title, context.content]
    if context.chatmetadata:
        context_parts.append(str(context.chatmetadata))

    context_text = "\n".join(context_parts)
    
    embedding = embed_text(context_text)
    
    context.embedding = embedding
    db.commit()
    
    return embedding

def generate_chat_message_embedding(db, chatmessage_id):
    
    message = db.query(ChatMessage).filter(ChatMessage.messageid == chatmessage_id).first()
    
    if not message:
        return None
    
    embedding = embed_text(message.content)
    
    message.embedding = embedding
    db.commit() 
    return embedding

def search_chat_contexts(db, user_id, query: str, limit: int = 5, similiarity_threshold: float = .5 ):
    
    query_embedding = embed_text(query)
    
    results = db.execute(
        select(
            ChatContext, (1 - ChatContext.embedding.cosine_distance(query_embedding)).label("similarity")
        ).where(
            ChatContext.createdby == user_id, ChatContext.isactive == True
        ).order_by(
            desc("similarity")
        ).where(
            (1 - ChatContext.embedding.cosine_distance(query_embedding)) >= similiarity_threshold
        ).limit(limit)
    ).all()
    
    formatted_results = [
        {
            "contextid": result.ChatContext.contextid,
            "content": result.ChatContext.content,
            "title": result.ChatContext.title,
            "author": result.ChatContext.createdby,
            "similarity": result.similarity
        }
        for result in results
    ]
    
    return formatted_results


    
    
    
    